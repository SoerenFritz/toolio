"""
Erstellt Word-Dokumente für die Lehrerwerkzeuge.
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt


def create_docx(
    title: str,
    content: str
) -> BytesIO:
    """
    Erstellt ein Word-Dokument und gibt es
    als BytesIO-Objekt zurück.
    """

    document = Document()

    heading = document.add_heading(title, level=1)
    heading.style.font.size = Pt(18)

    for line in content.splitlines():

        line = line.strip()

        if not line:
            document.add_paragraph("")
            continue

        document.add_paragraph(line)

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output